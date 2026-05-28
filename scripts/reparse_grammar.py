"""
剑桥英语语法 docx 重新解析器 — FINAL VERSION。

针对两本书（初级、中级），解析结构为：
  - 左手页 = 语法讲解，右手页 = 练习题
  - 用 'X.Y' 练习题号模式定位练习范围
  - 答案区在文档末尾用 'UNIT X' 标记
"""
import json, re, sys
from collections import OrderedDict

def clean(text):
    text = text.strip()
    return text if text and not text.isspace() else ''

def parse_book(paras_json_path, output_path, book_type='primary'):
    """
    book_type: 'primary' or 'intermediate'
    """
    with open(paras_json_path, 'r', encoding='utf-8') as f:
        paras = json.load(f)

    texts = {i: t for i, t in paras}
    n = max(texts.keys())

    # === Phase 1: 从目录提取所有单元标题 ===
    # TOC is in two ranges: 399-501 and 601-701
    toc_titles = {}
    for i, t in paras:
        t = t.strip()
        m = re.match(r'^(\d{1,3})\s{2,}(.+)$', t)
        if m:
            unit_num = int(m.group(1))
            if 1 <= unit_num <= 150:  # intermediate has up to 145
                title = m.group(2).strip()
                # Prefer titles from TOC ranges
                is_toc = (399 <= i <= 501) or (601 <= i <= 701)
                if unit_num not in toc_titles or is_toc:
                    toc_titles[unit_num] = title

    print(f'TOC titles: {len(toc_titles)}')

    # === Phase 2: 找到答案区 ===
    answer_start = n
    if book_type == 'primary':
        for i in sorted(texts.keys()):
            if i > 8000 and '练习答案' in texts[i].strip() and len(texts[i].strip()) < 20:
                answer_start = i
                break
    else:  # intermediate: answers start at "UNIT" markers before end
        # Search from p5000 onwards for first "UNIT X" marker in answers
        for i in sorted(texts.keys()):
            if i > 5000:
                t = texts[i].strip()
                m = re.match(r'^UNIT\s+(\d+)$', t)
                if m and i > len(paras) * 0.6:  # past ~60% of book
                    answer_start = i
                    break
    print(f'Answer section starts at paragraph {answer_start}')

    # === Phase 3: 找到主内容区 ===
    # Content starts at the first "Exercises Unit 1" marker
    content_start = 700
    for i in sorted(texts.keys()):
        if 700 < i < 2000:
            t = texts[i].strip()
            if re.match(r'Exercises\s+Unit\s+\d+', t):
                content_start = i
                break

    # Content ends before 补充练习/附录 or answer section
    # The real appendix section starts after p8000.
    # Earlier matches (like p1806 "→ 补充练习...") are cross-references, not section headers.
    content_end = answer_start
    for i in sorted(texts.keys()):
        if i > 8000 and i < answer_start - 500:
            t = texts[i].strip()
            # Standalone header: "附录N" or "补充练习"
            if (t.startswith('附录') and len(t) < 10) or (t == '补充练习'):
                content_end = i
                break

    print(f'Main content: paragraphs {content_start} - {content_end}')

    # === Phase 4: 找到所有 'X.Y' 练习标题 ===
    ex_headers = []
    for i in range(content_start, content_end):
        if i not in texts:
            continue
        t = texts[i].strip()
        m = re.match(r'^(\d{1,3})\.(\d{1,2})\s+(\S.*)', t)
        if m:
            unit = int(m.group(1))
            max_unit = 116 if book_type == 'primary' else 150
            if 1 <= unit <= max_unit:
                ex_headers.append((i, unit, m.group(2), m.group(3)))

    print(f'Exercise sub-headers found: {len(ex_headers)}')

    # === Phase 5: 每个单元的第一个练习簇 ===
    unit_ranges = {}
    for i, unit, ex_num, desc in ex_headers:
        if unit not in unit_ranges:
            unit_ranges[unit] = [i, i]
        else:
            gap = i - unit_ranges[unit][1]
            if gap <= 100:
                unit_ranges[unit][1] = i

    print(f'Units with exercise ranges: {len(unit_ranges)}')

    # === Phase 6: 构建单元内容 ===
    sorted_units = sorted(unit_ranges.keys(), key=lambda u: unit_ranges[u][0])
    units = OrderedDict()

    for i, unit_num in enumerate(sorted_units):
        ex_start, ex_end = unit_ranges[unit_num]

        # Explanation start: after previous unit's exercises
        if i == 0:
            expl_start = content_start
        else:
            prev_unit = sorted_units[i-1]
            expl_start = unit_ranges[prev_unit][1] + 1

        expl_end = ex_start

        # Title from TOC
        title = toc_titles.get(unit_num, '')

        # Collect explanation
        explanation = []
        for j in range(expl_start, expl_end):
            if j not in texts:
                continue
            t = clean(texts[j])
            if not t or len(t) < 2:
                continue
            if re.match(r'^\d{1,3}\.\d{1,2}\s', t):
                continue
            if re.match(r'Exercises\s+Unit', t):
                continue
            if '→Unit' in t or re.match(r'^\d{1,4}$', t):
                continue
            explanation.append(t)

        # Collect exercises
        if i + 1 < len(sorted_units):
            next_unit = sorted_units[i+1]
            ex_content_end = unit_ranges[next_unit][0]
        else:
            ex_content_end = content_end

        exercises = []
        for j in range(ex_start, min(ex_content_end, content_end)):
            if j not in texts:
                continue
            t = clean(texts[j])
            if not t:
                continue
            m = re.match(r'^(\d{1,3})\s{2,}(.+)$', t)
            if m and int(m.group(1)) in sorted_units and int(m.group(1)) > unit_num:
                break
            exercises.append(t)

        units[str(unit_num).zfill(3)] = {
            'title': title,
            'explanation': '\n'.join(explanation),
            'exercises': '\n'.join(exercises),
            'answers': {}
        }

    # === Phase 7: 解析答案 ===
    print(f'\nParsing answers from paragraph {answer_start}...')
    current_unit = None
    current_exercise = None
    answers = OrderedDict()

    for idx in range(answer_start, n):
        if idx not in texts:
            continue
        text = texts[idx].strip()
        if not text or len(text) < 2:
            continue

        # 'UNIT X' header (standalone or combined with '练习答案')
        m_unit = re.match(r'^UNIT\s+(\d{1,3})$', text)
        m_combined = re.match(r'^练习答案\s+UNIT\s+(\d+)', text)

        if m_unit:
            current_unit = int(m_unit.group(1))
            if current_unit not in answers:
                answers[current_unit] = OrderedDict()
            current_exercise = None
            continue

        if m_combined:
            current_unit = int(m_combined.group(1))
            if current_unit not in answers:
                answers[current_unit] = OrderedDict()
            current_exercise = None
            continue

        if current_unit is None:
            continue

        # Skip pure numbers (page numbers)
        if re.match(r'^\d{1,4}$', text):
            continue

        # Skip section headers without exercise numbers
        # Skip lines that are clearly not answers (like standalone '补充练习答案')
        if '补充练习答案' in text or '学习指导答案' in text:
            continue

        # Exercise sub-header: 'X.Y' or 'X. Y'
        m_ex = re.match(r'^(\d{1,3})\.\s*(\d{1,2})\b\s*(.*)', text)
        if m_ex:
            ex_unit = int(m_ex.group(1))
            ex_num = m_ex.group(2)
            rest = m_ex.group(3).strip()

            # Accept if unit matches or is nearby (handle missing UNIT headers)
            if current_unit and abs(ex_unit - current_unit) <= 2:
                if ex_unit != current_unit:
                    current_unit = ex_unit
                    if current_unit not in answers:
                        answers[current_unit] = OrderedDict()
                current_exercise = ex_num
                if current_unit not in answers:
                    answers[current_unit] = OrderedDict()
                if current_exercise not in answers[current_unit]:
                    answers[current_unit][current_exercise] = []
                if rest:
                    answers[current_unit][current_exercise].append(rest)
                continue

        # Regular answer line
        if current_unit and current_exercise is not None:
            if current_unit not in answers:
                answers[current_unit] = OrderedDict()
            if current_exercise not in answers[current_unit]:
                answers[current_unit][current_exercise] = []
            answers[current_unit][current_exercise].append(text)

    print(f'Parsed answers for {len(answers)} units')

    # Assign answers to units
    for unit_str in units:
        unit_num = int(unit_str)
        if unit_num in answers:
            ans_out = OrderedDict()
            for ex_num in sorted(answers[unit_num].keys(),
                                key=lambda x: int(x) if x.isdigit() else x):
                ans_out[ex_num] = '\n'.join(answers[unit_num][ex_num])
            units[unit_str]['answers'] = ans_out

    # === Stats ===
    with_expl = sum(1 for u in units.values() if u['explanation'].strip())
    with_ex = sum(1 for u in units.values() if u['exercises'].strip())
    with_ans = sum(1 for u in units.values() if u['answers'])
    print(f'Total units: {len(units)}')
    print(f'With explanation: {with_expl}')
    print(f'With exercises: {with_ex}')
    print(f'With answers: {with_ans}')

    # Write output
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(units, f, ensure_ascii=False, indent=2)
    print(f'Written to {output_path}')

    return units


def main():
    # === Primary Grammar ===
    print('='*60)
    print('PRIMARY GRAMMAR (Essential Grammar in Use)')
    print('='*60)
    units_p = parse_book(
        r'C:\Users\77630\.claude\skills\nce-coach\data\_primary_paras.json',
        r'C:\Users\77630\.claude\skills\nce-coach\data\grammar_primary_new.json',
        book_type='primary'
    )

    # Sample check
    for key in ['001', '005', '030', '045', '100']:
        if key in units_p:
            u = units_p[key]
            print(f'\nUnit {int(key)}: title={u["title"][:60]}')
            print(f'  explanation={len(u["explanation"])} chars, exercises={len(u["exercises"])} chars, answers={len(u["answers"])} groups')

    # === Intermediate Grammar ===
    print('\n' + '='*60)
    print('INTERMEDIATE GRAMMAR (English Grammar in Use)')
    print('='*60)

    # Load intermediate (should already be cached)
    intermediate_cache = r'C:\Users\77630\.claude\skills\nce-coach\data\_intermediate_paras.json'
    if not __import__('os').path.exists(intermediate_cache):
        from docx import Document
        doc = Document(r'D:\新概念\剑桥英语语法\剑桥中级英语语法_中文版.docx')
        print(f'Loaded intermediate: {len(doc.paragraphs)} paragraphs')
        paras = [(i, p.text) for i, p in enumerate(doc.paragraphs)]
        with open(intermediate_cache, 'w', encoding='utf-8') as f:
            json.dump(paras, f, ensure_ascii=False)
        print('Cached intermediate paragraphs')

    units_i = parse_book(
        intermediate_cache,
        r'C:\Users\77630\.claude\skills\nce-coach\data\grammar_intermediate_new.json',
        book_type='intermediate'
    )


if __name__ == '__main__':
    main()
